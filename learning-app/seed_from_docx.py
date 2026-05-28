"""
seed_from_docx.py
-----------------
Reads the outline table from a course notes .docx and creates
learning_notes_data.json pre-populated with all sections and topics.

Usage:
    python seed_from_docx.py 02_edukate_genai_llm_notes.docx \
        --title "Generative AI and Large Language Models" \
        --sub "Module 7 · L6 AI Engineer · Cambridge Spark"

The script reads the first table in the document. Expected columns:
  Col 0: Section heading (blank = use previous section)
  Col 1: Topic name
  Col 2: Covered marker (any non-empty value = done)
"""

import argparse, json, random, string, sys
from pathlib import Path


def uid():
    return "".join(random.choices(string.ascii_lowercase + string.digits, k=8))


def parse_table(docx_path):
    from docx import Document
    doc = Document(docx_path)
    if not doc.tables:
        print("No tables found in the document.")
        sys.exit(1)
    rows = []
    for i, row in enumerate(doc.tables[0].rows):
        if i == 0: continue
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2:
            rows.append(cells)
    return rows


def build_data(rows, title, sub):
    sections, order, current = {}, [], None
    for cells in rows:
        sec_raw   = cells[0].strip() if cells[0].strip() else None
        topic_raw = cells[1].strip() if len(cells) > 1 else ""
        covered   = cells[2].strip() if len(cells) > 2 else ""
        if not topic_raw: continue
        if sec_raw:
            current = sec_raw
            if current not in sections:
                sections[current] = {"id": uid(), "title": current, "topics": []}
                order.append(current)
        if current is None: current = "General"
        if current not in sections:
            sections[current] = {"id": uid(), "title": current, "topics": []}
            order.append(current)
        sections[current]["topics"].append({
            "id": uid(), "name": topic_raw.lstrip("* ").strip(),
            "done": bool(covered and covered not in ("-", "")),
            "notes_html": "", "resources": []
        })

    module = {"id": uid(), "title": title or "Course Module",
              "sections": [sections[s] for s in order]}
    return {"course_title": title or "Learning Notes", "course_sub": sub or "",
            "modules": [module]}


def main():
    p = argparse.ArgumentParser()
    p.add_argument("docx")
    p.add_argument("--title", default="")
    p.add_argument("--sub", default="")
    p.add_argument("--data", default="learning_notes_data.json")
    args = p.parse_args()

    if not Path(args.docx).exists():
        print(f"Error: {args.docx} not found"); sys.exit(1)

    try:
        from docx import Document
    except ImportError:
        print("Run: pip install python-docx --break-system-packages"); sys.exit(1)

    rows = parse_table(args.docx)
    print(f"Parsed {len(rows)} rows")

    data = build_data(rows, args.title, args.sub)
    total = sum(len(s["topics"]) for m in data["modules"] for s in m["sections"])
    done  = sum(1 for m in data["modules"] for s in m["sections"]
                for t in s["topics"] if t["done"])
    print(f"Topics: {total} ({done} pre-marked done)")

    if Path(args.data).exists():
        resp = input(f"{args.data} already exists. Overwrite? [y/N] ")
        if resp.lower() != "y":
            print("Aborted."); sys.exit(0)

    with open(args.data, "w") as f:
        json.dump(data, f, indent=2)
    print(f"Written to {args.data}")


if __name__ == "__main__":
    main()
